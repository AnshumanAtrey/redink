#!/usr/bin/env python3
"""
assemble_docx.py — merge enabled sections (per report-recipe.yaml) into the active framework's report.

Reads:
  - framework.yaml              — active framework name
  - report-recipe.yaml          — which sections to include (in order)
  - frameworks/<active>/manifest.yaml — framework metadata + section defaults
  - sections/<id>/section.md    — per-section frontmatter + template
  - jsons/*.json                — per-finding entries
  - brand-guidelines.md         — firm details
  - targets.yaml                — engagement scopes
  - engagement-summary.md       — exec summary narrative
  - templates/<framework>.docx (optional) — firm's styled template

Writes:
  - output/<framework>-report.docx (or .json for owasp-optrs)

Usage:
    python3 scripts/assemble_docx.py
    python3 scripts/assemble_docx.py --only FINDING-014
    python3 scripts/assemble_docx.py --no-images
    python3 scripts/assemble_docx.py --framework owasp-optrs

Requires: python-docx (for docx frameworks), pyyaml
    pip install python-docx pyyaml
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    import yaml  # type: ignore
except ImportError:
    sys.exit("error: pyyaml not installed. Run: pip install pyyaml")


REPO_ROOT = Path(__file__).resolve().parent.parent
SECTIONS_DIR = REPO_ROOT / "sections"
FRAMEWORKS_DIR = REPO_ROOT / "frameworks"
JSONS_DIR = REPO_ROOT / "jsons"
TEMPLATES_DIR = REPO_ROOT / "templates"
OUTPUT_DIR = REPO_ROOT / "output"
BRAND_FILE = REPO_ROOT / "brand-guidelines.md"
TARGETS_FILE = REPO_ROOT / "targets.yaml"
FRAMEWORK_FILE = REPO_ROOT / "framework.yaml"
RECIPE_FILE = REPO_ROOT / "report-recipe.yaml"
SUMMARY_FILE = REPO_ROOT / "engagement-summary.md"


def load_framework_name(override: str | None) -> str:
    if override:
        return override
    if not FRAMEWORK_FILE.exists():
        sys.exit(f"error: {FRAMEWORK_FILE} not found.")
    data = yaml.safe_load(FRAMEWORK_FILE.read_text())
    name = data.get("framework") if isinstance(data, dict) else None
    if not name:
        sys.exit("error: framework.yaml does not specify 'framework: <name>'.")
    return name


def load_framework_manifest(name: str) -> dict[str, Any]:
    fw_dir = FRAMEWORKS_DIR / name
    if not fw_dir.exists():
        avail = sorted(p.name for p in FRAMEWORKS_DIR.iterdir() if p.is_dir())
        sys.exit(f"error: framework '{name}' not found. Available: {', '.join(avail)}")
    manifest_path = fw_dir / "manifest.yaml"
    if not manifest_path.exists():
        sys.exit(f"error: {manifest_path} missing.")
    return yaml.safe_load(manifest_path.read_text())


def load_recipe() -> list[dict[str, Any]]:
    if not RECIPE_FILE.exists():
        sys.exit(f"error: {RECIPE_FILE} not found. Run /redink-recipe first.")
    data = yaml.safe_load(RECIPE_FILE.read_text())
    return [s for s in data.get("sections", []) if s.get("enabled")]


def parse_brand(path: Path) -> dict[str, Any]:
    if not path.exists():
        sys.exit(f"error: {path} not found.")
    text = path.read_text()
    if re.search(r"<[^>]+>", text):
        sys.exit("error: brand-guidelines.md still contains placeholders (<...>). Fill every field first.")
    brand: dict[str, Any] = {}
    table_row = re.compile(r"^\| ([^|]+?) \| ([^|]+?) \|", re.MULTILINE)
    for m in table_row.finditer(text):
        key, value = m.group(1).strip(), m.group(2).strip()
        if value and not key.startswith("-") and not value.startswith("-"):
            brand.setdefault(key, value)
    return brand


def parse_targets(path: Path) -> dict[str, Any]:
    if not path.exists():
        sys.exit(f"error: {path} not found.")
    return yaml.safe_load(path.read_text())


def parse_summary(path: Path) -> dict[str, str]:
    """Parse engagement-summary.md into {Overview, Summary of Findings, Key Observations}."""
    if not path.exists():
        sys.exit(f"error: {path} not found. Run /redink-build to generate template, then fill it in.")
    text = path.read_text()
    sections: dict[str, str] = {}
    current = None
    buf: list[str] = []
    for line in text.splitlines():
        m = re.match(r"^## (.+)$", line)
        if m:
            if current:
                sections[current] = "\n".join(buf).strip()
            current = m.group(1).strip()
            buf = []
        elif current:
            buf.append(line)
    if current:
        sections[current] = "\n".join(buf).strip()
    # Check for unfilled placeholders
    for k, v in sections.items():
        if re.search(r"<[^>]+>", v):
            sys.exit(f"error: engagement-summary.md '{k}' still contains placeholders. Fill it in first.")
    return sections


def load_findings(only: str | None) -> list[dict[str, Any]]:
    findings = []
    if not JSONS_DIR.exists():
        return findings
    for p in sorted(JSONS_DIR.glob("*.json")):
        if p.name in {"schema.json", "EXAMPLE.json"}:
            continue
        data = json.loads(p.read_text())
        fid = data.get("s_no") or data.get("id") or data.get("finding_id") or data.get("finding_ref")
        if only and fid != only:
            continue
        findings.append(data)
    if only:
        return findings
    return sorted(findings, key=lambda d: (d.get("s_no") or d.get("id") or d.get("finding_id") or d.get("finding_ref") or ""))


def parse_section_frontmatter(section_id: str) -> dict[str, Any]:
    md = SECTIONS_DIR / section_id / "section.md"
    if not md.exists():
        return {}
    text = md.read_text()
    m = re.match(r"^---\n(.*?)\n---", text, re.DOTALL)
    if not m:
        return {}
    return yaml.safe_load(m.group(1)) or {}


# ───── Section renderers ──────────────────────────────────────────────────────
# Each renderer takes the docx Document and the shared context, appends to the doc.

def render_cover_page(doc, ctx):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(ctx["manifest"].get("report_title", "VAPT Report"))
    run.bold = True
    run.font.size = Pt(26)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Submitted by: {ctx['brand'].get('Firm name', '<firm>')}").italic = True
    doc.add_paragraph()
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k in ["Engagement code", "Test duration (start)", "Test duration (end)", "Report version", "Submitted on"]:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = str(ctx["brand"].get(k, "N/A"))
    doc.add_page_break()


def render_executive_summary(doc, ctx):
    doc.add_heading("Executive Summary", level=1)
    summary = ctx["summary"]
    for sub in ["Overview", "Summary of Findings", "Key Observations & Critical Risks"]:
        if sub in summary:
            doc.add_heading(sub, level=2)
            doc.add_paragraph(summary[sub])


def render_scope_statement(doc, ctx):
    doc.add_heading("Engagement Scope", level=1)
    for scope in ctx["targets"].get("scopes", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{scope.get('name', '?')}: ").bold = True
        p.add_run(f"{scope.get('host', '?')} @ {scope.get('ip', '?')} — {', '.join(scope.get('services') or [])}")
    tw = ctx["targets"].get("testing_window", {})
    if tw:
        doc.add_paragraph(f"Testing window: {tw.get('start', '?')} to {tw.get('end', '?')} ({tw.get('timezone', '?')})")


def render_methodology(doc, ctx):
    doc.add_heading("Testing Methodology", level=1)
    method = ctx["brand"].get("Attack methodology narrative") or "See brand-guidelines.md §7."
    doc.add_paragraph(method)


def render_team_roster(doc, ctx):
    doc.add_heading("Technical Manpower", level=1)
    # Pull table from brand-guidelines.md §4 (already parsed as separate rows)
    doc.add_paragraph("See brand-guidelines.md §4 for the full team roster. The assembler will render the manpower table here in a future revision.")


def render_standards_adopted(doc, ctx):
    doc.add_heading("Standards, Frameworks, and Severity Scoring Systems", level=1)
    doc.add_paragraph("Standards adopted for this engagement (from brand-guidelines.md §5):")
    doc.add_paragraph("OWASP Top 10 · OWASP API Top 10 · OSSTMM v3 · SANS Top 25 · CIS Benchmarks · CVSS v3.1 · EPSS (FIRST.org)", style="Intense Quote")


def render_tools_used(doc, ctx):
    doc.add_heading("Tools, Scripts, and Frameworks Used", level=1)
    doc.add_paragraph("See brand-guidelines.md §6 for the firm's tool inventory.")


def render_summary_of_findings(doc, ctx):
    doc.add_heading("Summary of Findings", level=1)
    findings = ctx["findings"]
    counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0, "Informational": 0}
    for f in findings:
        sev = f.get("severity", "Informational")
        if sev in counts:
            counts[sev] += 1
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    head = table.add_row().cells
    head[0].text = "Severity"
    head[1].text = "Count"
    for sev, n in counts.items():
        row = table.add_row().cells
        row[0].text = sev
        row[1].text = str(n)
    total_row = table.add_row().cells
    total_row[0].text = "Total"
    total_row[1].text = str(sum(counts.values()))


def render_key_observations(doc, ctx):
    doc.add_heading("Key Observations & Critical Risks", level=1)
    if "Key Observations & Critical Risks" in ctx["summary"]:
        doc.add_paragraph(ctx["summary"]["Key Observations & Critical Risks"])


def render_aggregated_findings_table(doc, ctx):
    doc.add_heading("List of Vulnerable Locations Discovered", level=1)
    table = doc.add_table(rows=0, cols=4)
    table.style = "Light Grid Accent 1"
    head = table.add_row().cells
    head[0].text = "S.No."
    head[1].text = "Vulnerable Location / Parameter / Path"
    head[2].text = "Name of Vulnerability"
    head[3].text = "CVE / CWE"
    for f in ctx["findings"]:
        row = table.add_row().cells
        row[0].text = str(f.get("s_no") or f.get("id") or f.get("finding_id") or "?")
        row[1].text = f"{f.get('vulnerable_location', '')} · {f.get('vulnerable_path_port_url', '')} · {f.get('vulnerable_parameter') or '—'}"
        row[2].text = f.get("name_of_vulnerability") or f.get("title") or "Untitled"
        row[3].text = f.get("cve_cwe", "")


def render_per_finding_details(doc, ctx):
    from docx.shared import Inches, Pt, RGBColor
    doc.add_heading("Detailed Findings", level=1)
    severity_colors = {
        "Critical": RGBColor(0xC0, 0x00, 0x00),
        "High": RGBColor(0xE6, 0x6C, 0x00),
        "Medium": RGBColor(0xC9, 0xA2, 0x27),
        "Low": RGBColor(0x4C, 0x6E, 0xF5),
        "Informational": RGBColor(0x60, 0x60, 0x60),
    }
    for f in ctx["findings"]:
        fid = f.get("s_no") or f.get("id") or f.get("finding_id") or "?"
        title = f.get("name_of_vulnerability") or f.get("title") or "Untitled finding"
        severity = f.get("severity") or "Informational"

        header = doc.add_paragraph()
        sev_run = header.add_run(f"[{severity}] ")
        sev_run.bold = True
        sev_run.font.color.rgb = severity_colors.get(severity, RGBColor(0, 0, 0))
        name_run = header.add_run(f"{fid} — {title}")
        name_run.bold = True
        name_run.font.size = Pt(14)

        details = [
            ("S.No.", str(fid)),
            ("Name of Vulnerability", title),
            ("Severity", severity),
            ("Vulnerable Location", f.get("vulnerable_location", "")),
            ("Vulnerable Path/Port/URL", f.get("vulnerable_path_port_url", "")),
            ("Vulnerable Parameter", str(f.get("vulnerable_parameter") or "N/A")),
            ("CVE / CWE", f.get("cve_cwe", "")),
            ("CVSS / EPSS", f.get("cvss_epss_score", "")),
        ]
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for k, v in details:
            row = table.add_row().cells
            row[0].text = k
            row[1].text = v

        if f.get("description"):
            p = doc.add_paragraph()
            p.add_run("Description: ").bold = True
            p.add_run(f["description"])

        steps = f.get("poc_steps") or f.get("evidence") or []
        if steps:
            doc.add_paragraph().add_run("Proof of Concept").bold = True
            if f.get("poc_intro"):
                doc.add_paragraph(f["poc_intro"], style="Intense Quote")
            for step in steps:
                p = doc.add_paragraph(style="List Number")
                p.add_run(step.get("caption", ""))
                if ctx["include_images"]:
                    img_path = REPO_ROOT / step.get("image_path", "")
                    if img_path.exists():
                        try:
                            doc.add_picture(str(img_path), width=Inches(5.5))
                        except Exception as exc:
                            doc.add_paragraph(f"[image error: {img_path.name} — {exc}]")
                    else:
                        doc.add_paragraph(f"[missing image: {step.get('image_path', '?')}]")

        recs = f.get("recommendations") or []
        if recs:
            doc.add_paragraph().add_run("Recommendations").bold = True
            for r in recs:
                doc.add_paragraph(r, style="List Bullet")

        refs = f.get("references", "")
        if refs:
            doc.add_paragraph().add_run("References").bold = True
            if isinstance(refs, str):
                for line in refs.splitlines():
                    if line.strip():
                        doc.add_paragraph(line.strip(), style="List Bullet")

        if f.get("additional_observations"):
            doc.add_paragraph().add_run("Additional Observations").bold = True
            doc.add_paragraph(f["additional_observations"])

        # Audit footer (small, gray) — for evaluator re-verification
        audit = f.get("cwe_cve_audit", "")
        if audit:
            ap = doc.add_paragraph()
            run = ap.add_run(f"Source-audit: {audit}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            run.italic = True

        doc.add_page_break()


def render_limitations(doc, ctx):
    doc.add_heading("Limitations & Disclaimers", level=1)
    doc.add_paragraph(
        "Testing was conducted within the timebox stated in §4 (Engagement Scope). "
        "Findings reflect the system state during the testing window only. "
        "Subsequent changes may have introduced new vulnerabilities or remediated reported ones. "
        "The report represents the auditor's professional opinion based on evidence collected; "
        "it is not a guarantee of system security."
    )


def render_references(doc, ctx):
    doc.add_heading("References", level=1)
    sources = ctx["manifest"].get("official_sources") or []
    if sources:
        doc.add_paragraph().add_run("Framework standards").bold = True
        for s in sources:
            doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph().add_run("Vulnerability databases").bold = True
    for s in [
        "MITRE CWE — https://cwe.mitre.org/",
        "NIST NVD — https://nvd.nist.gov/",
        "FIRST.org CVSS — https://www.first.org/cvss/v3.1/specification-document",
        "FIRST.org EPSS — https://api.first.org/data/v1/epss",
    ]:
        doc.add_paragraph(s, style="List Bullet")


def render_placeholder(section_id, title):
    """Returns a renderer function that just adds a placeholder for sections without dedicated code yet."""
    def _r(doc, ctx):
        doc.add_heading(title, level=1)
        doc.add_paragraph(
            f"[Section {section_id}: see sections/{section_id}/section.md for the template. "
            f"This section's dedicated renderer is on the v0.2 roadmap.]"
        )
    return _r


SECTION_RENDERERS = {
    "01-cover-page": render_cover_page,
    "03-executive-summary": render_executive_summary,
    "04-scope-statement": render_scope_statement,
    "06-testing-methodology": render_methodology,
    "07-standards-adopted": render_standards_adopted,
    "08-team-roster": render_team_roster,
    "09-tools-used": render_tools_used,
    "11-summary-of-findings": render_summary_of_findings,
    "12-key-observations": render_key_observations,
    "13-aggregated-findings-table": render_aggregated_findings_table,
    "14-per-finding-details": render_per_finding_details,
    "21-limitations": render_limitations,
    "22-references": render_references,
}


def assemble_owasp_optrs(findings, brand, targets, manifest):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / "owasp-optrs-report.json"
    bundle = {
        "schema": "https://owasp.org/www-project-penetration-test-reporting-standard/",
        "engagement": {
            "firm": brand.get("Firm name", "N/A"),
            "engagement_code": brand.get("Engagement code", "N/A"),
            "test_start": brand.get("Test duration (start)", "N/A"),
            "test_end": brand.get("Test duration (end)", "N/A"),
        },
        "targets": targets.get("scopes", []),
        "findings": findings,
    }
    output.write_text(json.dumps(bundle, indent=2))
    return output


def assemble_docx(framework, manifest, recipe, brand, targets, summary, findings, include_images):
    try:
        from docx import Document  # type: ignore
    except ImportError:
        sys.exit("error: python-docx not installed. Run: pip install python-docx")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    template = TEMPLATES_DIR / f"{framework}.docx"
    doc = Document(str(template)) if template.exists() else Document()

    ctx = {
        "framework": framework,
        "manifest": manifest,
        "brand": brand,
        "targets": targets,
        "summary": summary,
        "findings": findings,
        "include_images": include_images,
    }

    for s in recipe:
        sid = s.get("id")
        if not sid:
            continue
        renderer = SECTION_RENDERERS.get(sid)
        if renderer:
            renderer(doc, ctx)
        else:
            # Placeholder for sections we haven't dedicated-rendered yet
            fm = parse_section_frontmatter(sid)
            title = fm.get("title", sid)
            render_placeholder(sid, title)(doc, ctx)

    output = OUTPUT_DIR / f"{framework}-report.docx"
    doc.save(str(output))
    return output


def main():
    parser = argparse.ArgumentParser(description="Assemble the active framework's report from sections + jsons + inputs.")
    parser.add_argument("--only", help="Only include this finding ID.", default=None)
    parser.add_argument("--no-images", action="store_true", help="Skip embedding screenshots.")
    parser.add_argument("--framework", help="Override framework.yaml.", default=None)
    args = parser.parse_args()

    framework = load_framework_name(args.framework)
    manifest = load_framework_manifest(framework)
    recipe = load_recipe()
    brand = parse_brand(BRAND_FILE)
    targets = parse_targets(TARGETS_FILE)
    summary = parse_summary(SUMMARY_FILE)
    findings = load_findings(args.only)
    if not findings:
        print("warning: no findings found in jsons/ — assembling shell report only", file=sys.stderr)

    if framework == "owasp-optrs":
        output = assemble_owasp_optrs(findings, brand, targets, manifest)
    else:
        output = assemble_docx(framework, manifest, recipe, brand, targets, summary, findings, include_images=not args.no_images)

    print(f"✓ Assembled ({framework}): {output.relative_to(REPO_ROOT)}")
    print(f"  Sections rendered: {len(recipe)}")
    print(f"  Findings included: {len(findings)}")


if __name__ == "__main__":
    main()
